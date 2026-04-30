"""Audit a batch-folder transcript run.

Checks for each source audio/video file in the folder:
  - Matching transcripts/<stem>_transcript.docx exists and is a reasonable size
  - Matching audio/<stem>.opus exists (warning, not error -- older runs
    may have skipped extraction)
  - Audio file size matches the bitrate * duration model (~24 kbps mono)
  - Docx transcript covers the full audio duration (no missing tail)
  - No large gaps in the timestamp sequence (no "skip in the middle")

Pure file-system inspection -- no DB or network. Safe to run on any
folder, past or future.
"""
from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import List, Optional

from docx import Document

from backend.app.services.notes.audio_probe import probe_duration_seconds


_AUDIO_EXT = {
    ".wav", ".mp3", ".m4a", ".opus", ".ogg", ".flac", ".aac", ".webm",
    ".mp4", ".mov", ".mkv", ".avi", ".m4v",
}
_VIDEO_EXT = {".mp4", ".mov", ".mkv", ".avi", ".m4v", ".webm"}

# Heuristics
_DOCX_MIN_KB                = 5      # smaller than this is suspicious
_OPUS_BITRATE_BPS_LONG      = 24_000   # bps used for >=40min audio
_OPUS_BITRATE_BPS_SHORT     = 48_000   # bps used for <40min audio
_OPUS_SIZE_TOLERANCE        = 0.5      # accept 0.5x to 2x of expected size
_TIMESTAMP_GAP_WARN_MINUTES = 5.0      # any gap > 5 min flagged
_TIMESTAMP_END_TOLERANCE_PCT = 0.10    # transcript should end within 10% of audio end
_WPM_WARN_LOW               = 30       # words/min below this is suspiciously sparse


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class FileAudit:
    name:               str
    is_video:           bool
    source_mb:          float
    duration_sec:       Optional[float]    # None if probe failed
    opus_exists:        bool
    opus_kb:            Optional[float]
    opus_size_ok:       Optional[bool]     # None if can't tell
    docx_exists:        bool
    docx_kb:            Optional[float]
    docx_size_ok:       Optional[bool]
    transcript_min_sec: Optional[float]    # earliest timestamp parsed
    transcript_max_sec: Optional[float]    # latest timestamp parsed
    transcript_gaps:    List[dict]         # [{start, end, gap_min}]
    transcript_coverage_pct: Optional[float]   # max_ts / duration
    word_count:         Optional[int]
    words_per_min:      Optional[float]
    issues:             List[str] = field(default_factory=list)


@dataclass
class AuditReport:
    folder:        str
    total_sources: int
    files:         List[FileAudit] = field(default_factory=list)
    summary:       dict            = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Audit logic
# ---------------------------------------------------------------------------

def audit_folder(folder_path: str) -> AuditReport:
    folder = Path(folder_path)
    if not folder.exists() or not folder.is_dir():
        raise FileNotFoundError(f"folder does not exist or is not a dir: {folder_path}")

    sources = sorted(
        [p for p in folder.iterdir()
         if p.is_file() and p.suffix.lower() in _AUDIO_EXT],
        key=lambda p: p.name.lower(),
    )

    transcripts_dir = folder / "transcripts"
    audio_dir       = folder / "audio"

    # Match the disambiguation rule from batch_scan: collisions on stem
    # produce {stem}_{ext}_transcript.docx, otherwise {stem}_transcript.docx.
    stem_collisions: Counter = Counter(p.stem for p in sources)

    audits: List[FileAudit] = []
    for src in sources:
        audits.append(_audit_one(src, transcripts_dir, audio_dir, stem_collisions))

    # Roll up summary counts
    n_complete   = sum(1 for f in audits if not f.issues)
    n_with_issue = sum(1 for f in audits if f.issues)
    n_missing_doc = sum(1 for f in audits if not f.docx_exists)
    n_missing_audio = sum(1 for f in audits if f.is_video and not f.opus_exists)
    n_with_gaps  = sum(1 for f in audits if f.transcript_gaps)

    return AuditReport(
        folder=str(folder),
        total_sources=len(sources),
        files=audits,
        summary={
            "complete":           n_complete,
            "with_issues":        n_with_issue,
            "missing_transcript": n_missing_doc,
            "missing_opus":       n_missing_audio,
            "with_coverage_gaps": n_with_gaps,
        },
    )


def _audit_one(
    src: Path,
    transcripts_dir: Path,
    audio_dir: Path,
    stem_collisions: Counter,
) -> FileAudit:
    issues: List[str] = []

    # Source size
    source_mb = round(src.stat().st_size / (1024 * 1024), 1)
    is_video  = src.suffix.lower() in _VIDEO_EXT

    # Probe duration (used as the reference for coverage / size sanity).
    try:
        duration_sec = probe_duration_seconds(str(src))
    except (RuntimeError, ValueError) as exc:
        duration_sec = None
        issues.append(f"could not probe duration: {exc}")

    # Disambiguate transcript filename if there's a stem collision.
    ext = src.suffix.lower()
    if stem_collisions[src.stem] > 1:
        transcript_name = f"{src.stem}_{ext.lstrip('.')}_transcript.docx"
    else:
        transcript_name = f"{src.stem}_transcript.docx"
    transcript_path = transcripts_dir / transcript_name
    opus_path       = audio_dir / f"{src.stem}.opus"

    # Audio (opus) presence + size sanity
    opus_exists = opus_path.exists()
    opus_kb     = round(opus_path.stat().st_size / 1024, 1) if opus_exists else None
    opus_size_ok: Optional[bool] = None
    if opus_exists and duration_sec is not None and duration_sec > 0:
        bitrate = _OPUS_BITRATE_BPS_LONG if duration_sec >= 40 * 60 else _OPUS_BITRATE_BPS_SHORT
        expected_kb = (bitrate * duration_sec / 8) / 1024
        ratio = (opus_kb or 0) / max(expected_kb, 1)
        opus_size_ok = (1 - _OPUS_SIZE_TOLERANCE) <= ratio <= (1 + 1.0)
        if not opus_size_ok:
            issues.append(
                f"opus size {opus_kb:.0f} KB seems off for {int(duration_sec/60)} min audio "
                f"(expected ~{expected_kb:.0f} KB, ratio {ratio:.2f})"
            )
    if is_video and not opus_exists:
        issues.append("video source has no extracted .opus in audio/")

    # Docx existence + size + content
    docx_exists = transcript_path.exists()
    docx_kb     = round(transcript_path.stat().st_size / 1024, 1) if docx_exists else None
    docx_size_ok: Optional[bool] = None
    transcript_min_sec: Optional[float] = None
    transcript_max_sec: Optional[float] = None
    transcript_gaps:    List[dict]     = []
    coverage_pct:       Optional[float] = None
    word_count:         Optional[int]   = None
    words_per_min:      Optional[float] = None

    if not docx_exists:
        issues.append("transcript .docx is missing")
    else:
        docx_size_ok = (docx_kb or 0) >= _DOCX_MIN_KB
        if not docx_size_ok:
            issues.append(f"docx is suspiciously small ({docx_kb} KB)")

        try:
            timestamps_sec, total_words = _read_docx_timestamps_and_words(transcript_path)
            word_count = total_words
            if timestamps_sec:
                transcript_min_sec = timestamps_sec[0]
                transcript_max_sec = timestamps_sec[-1]

                if duration_sec and duration_sec > 0:
                    coverage_pct = round(transcript_max_sec / duration_sec, 3)
                    end_gap_sec  = duration_sec - transcript_max_sec
                    end_gap_pct  = end_gap_sec / duration_sec
                    if end_gap_pct > _TIMESTAMP_END_TOLERANCE_PCT:
                        issues.append(
                            f"transcript ends at {_fmt_ts(transcript_max_sec)} but audio is "
                            f"{_fmt_ts(duration_sec)} -- last "
                            f"{_fmt_ts(end_gap_sec)} ({end_gap_pct:.0%}) appears missing"
                        )

                # Internal gaps
                for prev_sec, next_sec in zip(timestamps_sec, timestamps_sec[1:]):
                    gap_min = (next_sec - prev_sec) / 60.0
                    if gap_min > _TIMESTAMP_GAP_WARN_MINUTES:
                        transcript_gaps.append({
                            "start_ts": _fmt_ts(prev_sec),
                            "end_ts":   _fmt_ts(next_sec),
                            "gap_min":  round(gap_min, 1),
                        })
                if transcript_gaps:
                    summary = ", ".join(
                        f"{g['start_ts']} -> {g['end_ts']} ({g['gap_min']} min missing)"
                        for g in transcript_gaps
                    )
                    issues.append(f"timestamp gaps detected: {summary}")

                # Words-per-minute sanity
                if duration_sec and duration_sec > 0 and word_count is not None:
                    words_per_min = round(word_count / (duration_sec / 60.0), 1)
                    if words_per_min < _WPM_WARN_LOW:
                        issues.append(
                            f"sparse transcript: {word_count} words for {int(duration_sec/60)} "
                            f"min audio ({words_per_min} wpm; expected 80-200)"
                        )
            else:
                issues.append("docx exists but no [MM:SS] timestamps found -- can't verify coverage")
        except Exception as exc:    # noqa: BLE001
            issues.append(f"could not read docx for coverage check: {exc}")

    return FileAudit(
        name              = src.name,
        is_video          = is_video,
        source_mb         = source_mb,
        duration_sec      = duration_sec,
        opus_exists       = opus_exists,
        opus_kb           = opus_kb,
        opus_size_ok      = opus_size_ok,
        docx_exists       = docx_exists,
        docx_kb           = docx_kb,
        docx_size_ok      = docx_size_ok,
        transcript_min_sec= transcript_min_sec,
        transcript_max_sec= transcript_max_sec,
        transcript_gaps   = transcript_gaps,
        transcript_coverage_pct = coverage_pct,
        word_count        = word_count,
        words_per_min     = words_per_min,
        issues            = issues,
    )


# ---------------------------------------------------------------------------
# docx reading helpers
# ---------------------------------------------------------------------------

# Match [HH:MM:SS] or [MM:SS] anywhere in the text.
_TS_RX = re.compile(r"\[(\d{1,2}:\d{2}(?::\d{2})?)\]")


def _read_docx_timestamps_and_words(path: Path) -> tuple[list[float], int]:
    """Return (sorted list of timestamps in seconds, total word count) for a
    transcript docx. Handles both monolingual paragraph layout and bilingual
    table layout. Word count uses a naive whitespace split."""
    doc = Document(str(path))
    timestamps: list[float] = []
    total_words = 0

    # Paragraphs (monolingual + the H1/metadata + any review section)
    for para in doc.paragraphs:
        text = para.text or ""
        if not text:
            continue
        total_words += len(text.split())
        for m in _TS_RX.finditer(text):
            timestamps.append(_parse_ts(m.group(1)))

    # Tables (bilingual: time | original | translation)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text or ""
                if not text:
                    continue
                total_words += len(text.split())
                # In the bilingual table, col 0 is just "MM:SS" without brackets.
                stripped = text.strip()
                if re.fullmatch(r"\d{1,2}:\d{2}(?::\d{2})?", stripped):
                    timestamps.append(_parse_ts(stripped))
                else:
                    for m in _TS_RX.finditer(text):
                        timestamps.append(_parse_ts(m.group(1)))

    return sorted(timestamps), total_words


def _parse_ts(s: str) -> float:
    parts = [int(x) for x in s.split(":")]
    if len(parts) == 2:
        m, sec = parts
        return m * 60 + sec
    if len(parts) == 3:
        h, m, sec = parts
        return h * 3600 + m * 60 + sec
    return 0.0


def _fmt_ts(sec: Optional[float]) -> str:
    if sec is None:
        return "?"
    s = int(sec)
    if s < 60:
        return f"00:{s:02d}"
    if s < 3600:
        return f"{s // 60:02d}:{s % 60:02d}"
    return f"{s // 3600:d}:{(s % 3600) // 60:02d}:{s % 60:02d}"


# ---------------------------------------------------------------------------
# Markdown formatting
# ---------------------------------------------------------------------------

def format_audit_report(report: AuditReport) -> str:
    """Pretty-print an AuditReport as markdown."""
    lines: list[str] = []
    lines.append(f"# Batch audit: {report.folder}")
    lines.append("")
    s = report.summary
    lines.append(f"- Total source files: **{report.total_sources}**")
    lines.append(f"- Complete (no issues): **{s.get('complete', 0)}**")
    lines.append(f"- With issues: **{s.get('with_issues', 0)}**")
    if s.get("missing_transcript"):
        lines.append(f"  - Missing transcript .docx: {s['missing_transcript']}")
    if s.get("missing_opus"):
        lines.append(f"  - Video sources without extracted .opus: {s['missing_opus']}")
    if s.get("with_coverage_gaps"):
        lines.append(f"  - With internal coverage gaps (>5 min): {s['with_coverage_gaps']}")
    lines.append("")

    # Group: files with issues first, then OK files
    flagged = [f for f in report.files if f.issues]
    ok      = [f for f in report.files if not f.issues]

    if flagged:
        lines.append("## Files with issues")
        lines.append("")
        for f in flagged:
            lines.extend(_format_file(f, verbose=True))
            lines.append("")
    if ok:
        lines.append("## Clean files")
        lines.append("")
        for f in ok:
            lines.extend(_format_file(f, verbose=False))
        lines.append("")

    return "\n".join(lines)


def _format_file(f: FileAudit, *, verbose: bool) -> list[str]:
    out: list[str] = []
    icon = "WARN" if f.issues else "OK"
    out.append(f"### [{icon}] {f.name}")
    bits: list[str] = []
    bits.append(f"source {f.source_mb} MB")
    if f.duration_sec is not None:
        bits.append(f"duration {_fmt_ts(f.duration_sec)}")
    if f.opus_exists:
        bits.append(f"opus {f.opus_kb} KB")
    if f.docx_exists:
        bits.append(f"docx {f.docx_kb} KB")
    if f.word_count is not None:
        bits.append(f"{f.word_count} words")
    if f.words_per_min is not None:
        bits.append(f"{f.words_per_min} wpm")
    if f.transcript_coverage_pct is not None:
        bits.append(f"covers {f.transcript_coverage_pct:.0%} of audio")
    out.append("- " + " | ".join(bits))
    if f.issues:
        for issue in f.issues:
            out.append(f"  - **issue**: {issue}")
    return out


def report_to_dict(report: AuditReport) -> dict:
    return {
        "folder":        report.folder,
        "total_sources": report.total_sources,
        "files":         [asdict(f) for f in report.files],
        "summary":       report.summary,
    }
