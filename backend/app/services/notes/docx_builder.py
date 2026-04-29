"""Render a Note's polished transcript to a .docx file as bytes.

Extracted from backend.app.api.routers.v1.notes.export_note_as_docx so the
batch-folder transcription path can call it directly without going through
the HTTP layer.

The function is intentionally synchronous and takes a duck-typed `note`
that exposes:
  - title, note_id (str)
  - meeting_date (str | None)
  - company_tickers (list[str])
  - polished_transcript_meta (dict)

Raises ValueError when there are no segments to render.
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, Inches


def _render_markdown_into_doc(doc: "Document", md: str) -> None:
    """Render a small subset of markdown (headings, bullets, bold) into the
    document. Good enough for a Gemini-produced review -- not a full
    markdown engine. Anything we don't recognise is treated as a paragraph.
    """
    import re as _re
    # Split into lines once; keep blanks so we can detect paragraph breaks.
    paragraph_buf: list[str] = []

    def flush_paragraph():
        if paragraph_buf:
            text = " ".join(paragraph_buf).strip()
            if text:
                _add_inline_runs(doc.add_paragraph(), text)
            paragraph_buf.clear()

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            flush_paragraph()
            continue
        # Headings: ##, ###
        if line.startswith("### "):
            flush_paragraph()
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            flush_paragraph()
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            flush_paragraph()
            doc.add_heading(line[2:].strip(), level=1)
            continue
        # Bulleted list items
        m = _re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, m.group(1).strip())
            continue
        # Numbered list items
        m = _re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, m.group(1).strip())
            continue
        # Regular paragraph text -- accumulate and join on flush
        paragraph_buf.append(line)

    flush_paragraph()


def _add_inline_runs(paragraph, text: str) -> None:
    """Apply **bold** inline formatting to the runs of `paragraph`. Other
    inline markup (italic, code, links) is left as plain text."""
    import re as _re
    parts = _re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def build_note_docx(note: Any) -> bytes:
    meta = note.polished_transcript_meta or {}
    segments = list(meta.get("segments") or [])
    if not segments:
        raise ValueError("Note has no polished transcript segments to render.")

    is_bilingual = bool(meta.get("is_bilingual", False))
    language     = meta.get("language") or "en"
    audio_dur    = float(meta.get("audio_duration_sec") or 0.0)
    audio_min    = round(audio_dur / 60.0, 1)

    doc = Document()
    for section in doc.sections:
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    # Optional: AI-generated interview review at the top of the document.
    review_md = (meta.get("interview_review") or "").strip()
    if review_md:
        doc.add_heading("Interview Review (AI-generated)", level=1)
        _render_markdown_into_doc(doc, review_md)
        # Page break so the transcript starts on a fresh page.
        doc.add_page_break()

    doc.add_heading(note.title or f"Transcript {str(note.note_id)[:8]}", level=1)

    meta_bits: list[str] = []
    if note.meeting_date:
        meta_bits.append(str(note.meeting_date))
    if audio_dur > 0:
        meta_bits.append(f"audio {audio_min} min")
    meta_bits.append(f"language {language}{'/en' if is_bilingual else ''}")
    meta_bits.append(f"{len(segments)} segments")
    if note.company_tickers:
        meta_bits.append(", ".join(note.company_tickers))
    if meta_bits:
        para = doc.add_paragraph(" · ".join(meta_bits))   # · = middle dot, matches original endpoint output
        for run in para.runs:
            run.font.size = Pt(9)
            run.italic = True

    if is_bilingual:
        translation_label = meta.get("translation_label") or "English"
        table = doc.add_table(rows=1, cols=3)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass
        hdr = table.rows[0].cells
        hdr[0].text = "Time"
        hdr[1].text = "原文"   # matches original /export.docx output
        hdr[2].text = translation_label
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for seg in segments:
            row = table.add_row().cells
            row[0].text = (seg.get("timestamp") or "").strip()
            speaker = (seg.get("speaker") or "").strip()
            orig    = (seg.get("text_original") or "").strip()
            row[1].text = (f"[{speaker}] {orig}" if speaker else orig)
            row[2].text = (seg.get("text_english") or "").strip()
    else:
        for seg in segments:
            ts      = (seg.get("timestamp") or "").strip()
            speaker = (seg.get("speaker") or "").strip()
            text    = (seg.get("text_original") or "").strip()
            p = doc.add_paragraph()
            ts_run = p.add_run(f"[{ts}] " if ts else "")
            ts_run.bold = True
            if speaker:
                sp_run = p.add_run(f"{speaker}: ")
                sp_run.italic = True
            p.add_run(text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
