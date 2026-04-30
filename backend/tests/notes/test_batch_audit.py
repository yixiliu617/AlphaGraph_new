"""Unit tests for batch_audit.

Real ffmpeg/ffprobe calls are skipped via patching; tests focus on the
audit logic (matching, size checks, timestamp parsing, gap detection).
"""
import io
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from backend.app.services.notes.batch_audit import (
    audit_folder, format_audit_report, _parse_ts, _fmt_ts,
)


# ---------------------------------------------------------------------------
# helpers: build a transcript .docx with specified timestamps
# ---------------------------------------------------------------------------

def _make_transcript_docx(out_path: Path, timestamps_sec: list[int], words_per_seg: int = 8):
    """Build a docx with paragraphs of '[MM:SS] body' for each timestamp."""
    doc = Document()
    doc.add_heading("Test Note", level=1)
    for ts in timestamps_sec:
        m = ts // 60
        s = ts % 60
        p = doc.add_paragraph()
        p.add_run(f"[{m:02d}:{s:02d}] ").bold = True
        p.add_run("hello " * words_per_seg)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# _parse_ts / _fmt_ts
# ---------------------------------------------------------------------------

def test_parse_ts_mmss():
    assert _parse_ts("12:34") == 12 * 60 + 34


def test_parse_ts_hhmmss():
    assert _parse_ts("1:23:45") == 1 * 3600 + 23 * 60 + 45


def test_fmt_ts_short():
    assert _fmt_ts(45) == "00:45"


def test_fmt_ts_medium():
    assert _fmt_ts(3 * 60 + 15) == "03:15"


def test_fmt_ts_long():
    assert _fmt_ts(2 * 3600 + 5 * 60 + 7) == "2:05:07"


# ---------------------------------------------------------------------------
# audit_folder happy path / missing files
# ---------------------------------------------------------------------------

@pytest.fixture
def stub_probe_30min():
    with patch(
        "backend.app.services.notes.batch_audit.probe_duration_seconds",
        return_value=30 * 60.0,    # 30 minute audio
    ):
        yield


def test_clean_folder_passes(tmp_path: Path, stub_probe_30min):
    src = tmp_path / "interview.mp4"
    src.write_bytes(b"\0" * 100_000_000)   # 100 MB source

    # Reasonable opus: 30 min * 48kbps = ~10.8 MB. We use 9 MB to be in tolerance.
    opus = tmp_path / "audio" / "interview.opus"
    opus.parent.mkdir()
    opus.write_bytes(b"\0" * 9_000_000)

    # Reasonable docx with timestamps every 30s for the full 30 min, no gaps.
    transcript = tmp_path / "transcripts" / "interview_transcript.docx"
    timestamps = list(range(0, 30 * 60, 30))   # 0:00, 0:30, ..., 29:30
    _make_transcript_docx(transcript, timestamps, words_per_seg=20)

    report = audit_folder(str(tmp_path))
    assert report.total_sources == 1
    f = report.files[0]
    assert not f.issues, f"expected no issues, got {f.issues}"
    assert f.transcript_max_sec is not None and f.transcript_max_sec >= 29 * 60
    assert f.transcript_coverage_pct is not None and f.transcript_coverage_pct >= 0.9
    assert f.word_count is not None and f.word_count > 100


def test_missing_transcript(tmp_path: Path, stub_probe_30min):
    src = tmp_path / "missing.mp4"
    src.write_bytes(b"\0" * 1024)
    report = audit_folder(str(tmp_path))
    issues = report.files[0].issues
    assert any("missing" in i.lower() for i in issues)


def test_missing_opus_for_video(tmp_path: Path, stub_probe_30min):
    src = tmp_path / "v.mp4"
    src.write_bytes(b"\0" * 1024)
    transcript = tmp_path / "transcripts" / "v_transcript.docx"
    _make_transcript_docx(transcript, [0, 30, 60], words_per_seg=20)

    report = audit_folder(str(tmp_path))
    issues = report.files[0].issues
    assert any("opus" in i.lower() for i in issues)


def test_no_opus_warning_for_audio_source(tmp_path: Path, stub_probe_30min):
    """For audio sources (mp3/wav/etc) we don't warn about missing opus --
    only video sources need extraction."""
    src = tmp_path / "audio_only.mp3"
    src.write_bytes(b"\0" * 1024)
    transcript = tmp_path / "transcripts" / "audio_only_transcript.docx"
    timestamps = list(range(0, 30 * 60, 30))
    _make_transcript_docx(transcript, timestamps, words_per_seg=20)

    report = audit_folder(str(tmp_path))
    issues = report.files[0].issues
    assert not any("opus" in i.lower() for i in issues), f"unexpected opus warning: {issues}"


# ---------------------------------------------------------------------------
# coverage gaps + truncated transcript
# ---------------------------------------------------------------------------

def test_internal_gap_flagged(tmp_path: Path, stub_probe_30min):
    """A 10-minute hole between consecutive timestamps should be flagged."""
    src = tmp_path / "gap.mp4"
    src.write_bytes(b"\0" * 1024)
    opus = tmp_path / "audio" / "gap.opus"
    opus.parent.mkdir()
    opus.write_bytes(b"\0" * 9_000_000)

    # Timestamps: 0..5 min, then jump to 15..30 min. 10-min gap from 5 to 15.
    timestamps = list(range(0, 5 * 60, 30)) + list(range(15 * 60, 30 * 60, 30))
    transcript = tmp_path / "transcripts" / "gap_transcript.docx"
    _make_transcript_docx(transcript, timestamps, words_per_seg=20)

    report = audit_folder(str(tmp_path))
    f = report.files[0]
    assert f.transcript_gaps, f"expected at least one gap to be detected"
    gap = f.transcript_gaps[0]
    assert gap["gap_min"] >= 9.0    # ~10 min gap
    assert any("gap" in i.lower() for i in f.issues)


def test_truncated_tail_flagged(tmp_path: Path, stub_probe_30min):
    """Transcript stopping at 15:00 for a 30-min audio = 50% missing. Flagged."""
    src = tmp_path / "truncated.mp4"
    src.write_bytes(b"\0" * 1024)
    opus = tmp_path / "audio" / "truncated.opus"
    opus.parent.mkdir()
    opus.write_bytes(b"\0" * 9_000_000)

    timestamps = list(range(0, 15 * 60, 30))   # only first 15 min covered
    transcript = tmp_path / "transcripts" / "truncated_transcript.docx"
    _make_transcript_docx(transcript, timestamps, words_per_seg=20)

    report = audit_folder(str(tmp_path))
    f = report.files[0]
    assert f.transcript_max_sec is not None
    assert f.transcript_max_sec < 16 * 60      # ends around 15 min
    assert any("missing" in i.lower() and "audio" in i.lower() for i in f.issues), (
        f"expected truncated-tail warning, got issues: {f.issues}"
    )


# ---------------------------------------------------------------------------
# size sanity checks
# ---------------------------------------------------------------------------

def test_opus_too_small_flagged(tmp_path: Path, stub_probe_30min):
    """An opus file way smaller than expected for the audio duration is suspicious."""
    src = tmp_path / "small.mp4"
    src.write_bytes(b"\0" * 1024)
    opus = tmp_path / "audio" / "small.opus"
    opus.parent.mkdir()
    opus.write_bytes(b"\0" * 50_000)    # 50 KB for 30 min (expected ~10 MB)

    transcript = tmp_path / "transcripts" / "small_transcript.docx"
    _make_transcript_docx(transcript, list(range(0, 30 * 60, 30)), words_per_seg=20)

    report = audit_folder(str(tmp_path))
    f = report.files[0]
    assert any("opus size" in i.lower() for i in f.issues), f"expected opus size warning, got {f.issues}"


def test_docx_too_small_flagged(tmp_path: Path, stub_probe_30min):
    """A nearly-empty docx (under 5KB) is suspicious."""
    src = tmp_path / "tiny.mp4"
    src.write_bytes(b"\0" * 1024)
    opus = tmp_path / "audio" / "tiny.opus"
    opus.parent.mkdir()
    opus.write_bytes(b"\0" * 9_000_000)
    # 1-segment transcript = very small docx
    transcript = tmp_path / "transcripts" / "tiny_transcript.docx"
    _make_transcript_docx(transcript, [0], words_per_seg=2)

    report = audit_folder(str(tmp_path))
    f = report.files[0]
    # Either size warning OR sparse-transcript warning (both legit).
    assert any(
        ("docx" in i.lower() and "small" in i.lower()) or
        ("sparse" in i.lower())
        for i in f.issues
    ), f"expected size/sparse warning, got {f.issues}"


# ---------------------------------------------------------------------------
# format_audit_report (smoke)
# ---------------------------------------------------------------------------

def test_format_audit_report_smoke(tmp_path: Path, stub_probe_30min):
    src = tmp_path / "x.mp4"
    src.write_bytes(b"\0" * 1024)
    report = audit_folder(str(tmp_path))
    md = format_audit_report(report)
    assert "Batch audit" in md
    assert str(tmp_path) in md
    assert "x.mp4" in md
