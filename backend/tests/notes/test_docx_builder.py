"""Unit tests for the extracted docx-building helper."""
import io

import pytest
from docx import Document

from backend.app.services.notes.docx_builder import build_note_docx


class _FakeNote:
    """Minimal duck-typed stand-in for the ORM Note object the helper reads."""
    def __init__(
        self,
        *,
        note_id="abc12345-def6-7890",
        title="My Test Note",
        meeting_date=None,
        company_tickers=None,
        polished_transcript_meta=None,
    ):
        self.note_id = note_id
        self.title = title
        self.meeting_date = meeting_date
        self.company_tickers = company_tickers or []
        self.polished_transcript_meta = polished_transcript_meta or {}


def _read_docx(buf: bytes) -> Document:
    return Document(io.BytesIO(buf))


def test_returns_bytes_of_valid_docx():
    note = _FakeNote(polished_transcript_meta={
        "language": "en",
        "is_bilingual": False,
        "segments": [
            {"timestamp": "00:01", "speaker": "Alice", "text_original": "Hello world."},
        ],
        "audio_duration_sec": 60.0,
    })
    out = build_note_docx(note)
    assert isinstance(out, bytes)
    assert len(out) > 1000  # any real .docx is at least 1 KB
    doc = _read_docx(out)
    paragraphs = [p.text for p in doc.paragraphs]
    assert any("My Test Note" in p for p in paragraphs)


def test_monolingual_renders_paragraphs_not_table():
    note = _FakeNote(polished_transcript_meta={
        "language": "en",
        "is_bilingual": False,
        "segments": [
            {"timestamp": "00:01", "speaker": "Alice", "text_original": "Hello."},
            {"timestamp": "00:05", "speaker": "Bob",   "text_original": "Goodbye."},
        ],
    })
    doc = _read_docx(build_note_docx(note))
    assert len(doc.tables) == 0
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "Hello." in body
    assert "Goodbye." in body


def test_bilingual_renders_three_column_table():
    note = _FakeNote(polished_transcript_meta={
        "language": "zh",
        "is_bilingual": True,
        "translation_label": "English",
        "segments": [
            {"timestamp": "00:01", "speaker": "A", "text_original": "Ni hao",
             "text_english": "Hello"},
            {"timestamp": "00:05", "speaker": "B", "text_original": "Zai jian",
             "text_english": "Goodbye"},
        ],
    })
    doc = _read_docx(build_note_docx(note))
    assert len(doc.tables) == 1
    rows = doc.tables[0].rows
    assert len(rows) == 3  # 1 header + 2 segments
    assert rows[0].cells[2].text == "English"
    assert rows[1].cells[2].text == "Hello"
    assert rows[2].cells[2].text == "Goodbye"


def test_empty_segments_raises():
    note = _FakeNote(polished_transcript_meta={"segments": []})
    with pytest.raises(ValueError, match="no polished transcript segments"):
        build_note_docx(note)


def test_interview_review_renders_at_top():
    """When meta has interview_review markdown, it appears as a labelled
    section before the regular title + transcript content."""
    note = _FakeNote(
        title="Candidate XYZ",
        polished_transcript_meta={
            "language": "en",
            "is_bilingual": False,
            "segments": [
                {"timestamp": "00:01", "speaker": "interviewer", "text_original": "Hi."},
            ],
            "interview_review": (
                "## Interviewee assessment\n"
                "- **Strengths**: clear on Python.\n"
                "- **Weaknesses**: vague on systems design.\n\n"
                "## Interviewer review\n"
                "- Question on REST APIs worked well.\n"
            ),
        },
    )
    doc = _read_docx(build_note_docx(note))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    # The review heading should come before the transcript title heading.
    review_idx = next((i for i, h in enumerate(headings) if "Interview Review" in h), -1)
    title_idx  = next((i for i, h in enumerate(headings) if "Candidate XYZ" in h), -1)
    assert review_idx >= 0, f"missing review heading in {headings}"
    assert title_idx  >  review_idx, f"review must precede title (got idx {review_idx}, {title_idx})"
    # Section subheadings render too
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "Interviewee assessment" in body
    assert "Interviewer review" in body
    assert "Strengths" in body  # bold tags get rendered as plain text in this minimal renderer


def test_no_interview_review_means_no_review_section():
    """When interview_review is empty/missing, the review heading is NOT added."""
    note = _FakeNote(
        title="Plain note",
        polished_transcript_meta={
            "language": "en",
            "is_bilingual": False,
            "segments": [{"timestamp": "00:01", "speaker": "", "text_original": "x"}],
        },
    )
    doc = _read_docx(build_note_docx(note))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert not any("Interview Review" in h for h in headings), f"unexpected review heading in {headings}"
